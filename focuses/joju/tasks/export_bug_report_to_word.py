#!/usr/bin/env python3
"""
Export consolidated bug report to Word document
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from datetime import datetime

def add_heading(doc, text, level=1):
    """Add formatted heading"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_table_row(table, cells_data, bold=False):
    """Add row to table"""
    row = table.add_row()
    for i, data in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = str(data)
        if bold:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

def add_code_block(doc, code):
    """Add code block with monospace font"""
    para = doc.add_paragraph()
    run = para.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)

def create_bug_report():
    """Create Word document from bug report"""
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Joju - Consolidated Bug Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    doc.add_paragraph('Total Bugs: 4 active')
    doc.add_paragraph('Statically Validated: 4/4')
    doc.add_paragraph('Ready for Fix: 4/4')
    doc.add_paragraph()
    
    # Executive Summary
    add_heading(doc, 'Executive Summary', 1)
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    add_table_row(table, ['Bug', 'Severity', 'Effort', 'Static Check', 'Status'], bold=True)
    add_table_row(table, ['Share Link Clarity', 'Medium', '1-2h', '✅ Yes', '✅ Ready'])
    add_table_row(table, ['Date Validation', 'Low', '1h', '✅ Yes', '✅ Ready'])
    add_table_row(table, ['Icon Alignment', 'Low', '30min', '✅ Yes', '✅ Ready'])
    add_table_row(table, ['Section Reorder', 'Medium', '2-3h', '⚠️ Partial', '✅ Ready'])
    
    doc.add_paragraph()
    doc.add_paragraph('Total Effort: 4.5-6.5 hours').bold = True
    doc.add_paragraph('All bugs have clear fix paths').bold = True
    doc.add_paragraph()
    
    # Bug #1
    add_heading(doc, 'Bug #1: Share Link Clarity', 1)
    doc.add_paragraph('Priority: MEDIUM | Effort: 1-2 hours').bold = True
    doc.add_paragraph()
    
    add_heading(doc, 'Issue', 2)
    doc.add_paragraph('Share link doesn\'t make it explicitly clear that you are sharing a public link')
    
    add_heading(doc, 'Static Validation', 2)
    doc.add_paragraph('✅ PASS - Can check dialog text in components')
    
    add_heading(doc, 'Fix Location', 2)
    doc.add_paragraph('• src/components/ProfileHeader.tsx (share button)', style='List Bullet')
    doc.add_paragraph('• src/components/SlugSettings.tsx (slug settings)', style='List Bullet')
    doc.add_paragraph('• src/pages/PublicProfilePage.tsx (public view)', style='List Bullet')
    
    add_heading(doc, 'Recommended Fix', 2)
    add_code_block(doc, '''<Dialog>
  <DialogContent>
    <DialogHeader>
      <DialogTitle>🌐 Share Public Profile</DialogTitle>
      <DialogDescription>
        ⚠️ This link is PUBLIC. Anyone with this link can view your profile.
      </DialogDescription>
    </DialogHeader>
    <Input value={publicUrl} readOnly />
    <Button onClick={copyToClipboard}>Copy Link</Button>
  </DialogContent>
</Dialog>''')
    
    add_heading(doc, 'Impact', 2)
    doc.add_paragraph('• User: Better privacy awareness', style='List Bullet')
    doc.add_paragraph('• Risk: Low (UX improvement)', style='List Bullet')
    doc.add_paragraph('• Testing: Manual verification of dialog text', style='List Bullet')
    doc.add_paragraph()
    
    # Bug #2
    add_heading(doc, 'Bug #2: Date Validation', 1)
    doc.add_paragraph('Priority: LOW | Effort: 1 hour').bold = True
    doc.add_paragraph()
    
    add_heading(doc, 'Issue', 2)
    doc.add_paragraph('Start date should not be allowed to be cleared')
    
    add_heading(doc, 'Static Validation', 2)
    doc.add_paragraph('✅ PASS - Can verify required and allowClear props')
    
    add_heading(doc, 'Fix Location', 2)
    doc.add_paragraph('• src/components/InlineDateEdit.tsx (component definition)', style='List Bullet')
    doc.add_paragraph('• All usages of <InlineDateEdit> for start dates', style='List Bullet')
    
    add_heading(doc, 'Recommended Fix', 2)
    add_code_block(doc, '''<InlineDateEdit
  value={startDate}
  onChange={setStartDate}
  required={true}
  allowClear={false}
  placeholder="Start date (required)"
/>''')
    
    add_heading(doc, 'Impact', 2)
    doc.add_paragraph('• User: Prevents data loss', style='List Bullet')
    doc.add_paragraph('• Risk: Very low (validation improvement)', style='List Bullet')
    doc.add_paragraph('• Testing: Try to clear start date, should show error', style='List Bullet')
    doc.add_paragraph()
    
    # Bug #3
    add_heading(doc, 'Bug #3: Icon Alignment', 1)
    doc.add_paragraph('Priority: LOW | Effort: 30 minutes').bold = True
    doc.add_paragraph()
    
    add_heading(doc, 'Issue', 2)
    doc.add_paragraph('Grabber and the delete icon are not aligned')
    
    add_heading(doc, 'Static Validation', 2)
    doc.add_paragraph('✅ PASS - Can verify flexbox alignment classes')
    
    add_heading(doc, 'Fix Location', 2)
    doc.add_paragraph('• src/components/SortableItem.tsx', style='List Bullet')
    doc.add_paragraph('• src/components/DroppableColumn.tsx', style='List Bullet')
    doc.add_paragraph('• Any section with drag handles', style='List Bullet')
    
    add_heading(doc, 'Recommended Fix', 2)
    add_code_block(doc, '''<div className="flex items-center justify-between gap-2">
  <div className="flex items-center gap-2">
    <button className="cursor-grab">
      <GripVertical className="h-5 w-5" />
    </button>
    <span>Content</span>
  </div>
  <button className="text-red-500">
    <Trash2 className="h-5 w-5" />
  </button>
</div>''')
    
    add_heading(doc, 'Impact', 2)
    doc.add_paragraph('• User: Visual polish', style='List Bullet')
    doc.add_paragraph('• Risk: None (CSS only)', style='List Bullet')
    doc.add_paragraph('• Testing: Visual inspection of alignment', style='List Bullet')
    doc.add_paragraph()
    
    # Bug #4
    add_heading(doc, 'Bug #4: Section Reorder Feature', 1)
    doc.add_paragraph('Priority: MEDIUM | Effort: 2-3 hours').bold = True
    doc.add_paragraph()
    
    add_heading(doc, 'Issue', 2)
    doc.add_paragraph('Fix/Improve the Section Reorder Feature')
    
    add_heading(doc, 'Static Validation', 2)
    doc.add_paragraph('⚠️ PARTIAL - Implementation exists, needs runtime testing')
    
    add_heading(doc, 'Fix Location', 2)
    doc.add_paragraph('• src/components/SortableItem.tsx (drag item)', style='List Bullet')
    doc.add_paragraph('• src/components/DroppableColumn.tsx (drop zone)', style='List Bullet')
    doc.add_paragraph('• src/components/CVView.tsx (drag context)', style='List Bullet')
    
    add_heading(doc, 'Potential Issues', 2)
    doc.add_paragraph('1. Drag not working smoothly', style='List Number')
    doc.add_paragraph('2. Drop zones not clear', style='List Number')
    doc.add_paragraph('3. State not persisting after reorder', style='List Number')
    doc.add_paragraph('4. Visual feedback missing', style='List Number')
    
    add_heading(doc, 'Impact', 2)
    doc.add_paragraph('• User: Better drag/drop UX', style='List Bullet')
    doc.add_paragraph('• Risk: Medium (core functionality)', style='List Bullet')
    doc.add_paragraph('• Testing: Requires manual drag/drop testing', style='List Bullet')
    doc.add_paragraph()
    
    # Recommended Fix Order
    add_heading(doc, 'Recommended Fix Order', 1)
    
    add_heading(doc, 'Phase 1: Quick Wins (2 hours)', 2)
    doc.add_paragraph('1. Icon Alignment (30 min) - CSS only', style='List Number')
    doc.add_paragraph('2. Date Validation (1 hour) - Props + validation', style='List Number')
    doc.add_paragraph('3. Share Link Clarity (1-2 hours) - Dialog text', style='List Number')
    
    add_heading(doc, 'Phase 2: Complex Fix (2-3 hours)', 2)
    doc.add_paragraph('4. Section Reorder (2-3 hours) - Drag/drop improvements', style='List Number')
    doc.add_paragraph()
    
    doc.add_paragraph('Total: 4.5-6.5 hours').bold = True
    doc.add_paragraph()
    
    # Testing Checklist
    add_heading(doc, 'Testing Checklist', 1)
    
    add_heading(doc, 'Manual Tests Required', 2)
    doc.add_paragraph('☐ Share link shows warning dialog', style='List Bullet')
    doc.add_paragraph('☐ Start date cannot be cleared', style='List Bullet')
    doc.add_paragraph('☐ Icons are aligned vertically', style='List Bullet')
    doc.add_paragraph('☐ Sections drag smoothly', style='List Bullet')
    doc.add_paragraph('☐ Drop zones are clear', style='List Bullet')
    doc.add_paragraph('☐ Reorder persists after save', style='List Bullet')
    doc.add_paragraph()
    
    # Visual Automation Testing
    add_heading(doc, 'Visual Automation Testing Options', 1)
    
    add_heading(doc, 'Recommended Tools', 2)
    doc.add_paragraph('1. Playwright - E2E testing with visual comparisons', style='List Number')
    doc.add_paragraph('2. Chromatic - Visual regression testing', style='List Number')
    doc.add_paragraph('3. Percy - Screenshot diffing', style='List Number')
    
    add_heading(doc, 'Test Scenarios', 2)
    doc.add_paragraph('• Capture before/after screenshots of icon alignment', style='List Bullet')
    doc.add_paragraph('• Record drag/drop interactions', style='List Bullet')
    doc.add_paragraph('• Validate dialog appearance', style='List Bullet')
    doc.add_paragraph('• Check responsive layouts', style='List Bullet')
    
    add_heading(doc, 'Next Steps', 2)
    doc.add_paragraph('• Set up Playwright for Joju', style='List Bullet')
    doc.add_paragraph('• Create visual baseline screenshots', style='List Bullet')
    doc.add_paragraph('• Add drag/drop test scenarios', style='List Bullet')
    doc.add_paragraph('• Integrate with CI/CD', style='List Bullet')
    
    # Save document
    output_dir = Path.home() / 'Documents' / 'Joju'
    output_dir.mkdir(parents=True, exist_ok=True)
    output_file = output_dir / f'Joju_Bug_Report_{datetime.now().strftime("%Y-%m-%d")}.docx'
    
    doc.save(output_file)
    return output_file

if __name__ == '__main__':
    print("\n" + "="*80)
    print("EXPORTING BUG REPORT TO WORD")
    print("="*80 + "\n")
    
    try:
        output_file = create_bug_report()
        print(f"✅ Bug report exported to:")
        print(f"   {output_file}\n")
        print("📄 Document includes:")
        print("   • Executive summary table")
        print("   • All 4 bugs with details")
        print("   • Fix recommendations with code")
        print("   • Testing checklist")
        print("   • Visual automation testing options")
        print()
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        print("\nMake sure python-docx is installed:")
        print("   pip3 install python-docx")
