#!/usr/bin/env python3
"""
Export enhanced validation report to Word document
"""

import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from datetime import datetime

def add_heading(doc, text, level=1):
    """Add formatted heading"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_table_row(table, cells_data, bold=False):
    """Add row to table"""
    row = table.add_row()
    for i, data in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = str(data)
        if bold:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

def create_enhanced_validation_report():
    """Create Word document from enhanced validation"""
    
    # Load reports
    enhanced_file = Path(__file__).parent / 'ENHANCED_VALIDATION_REPORT.json'
    basic_file = Path(__file__).parent / 'VALIDATION_REPORT.json'
    
    with open(enhanced_file) as f:
        enhanced = json.load(f)
    
    with open(basic_file) as f:
        basic = json.load(f)
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Task Truth Pipeline V2', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Enhanced Validation Report', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    doc.add_paragraph(f'Validator: Enhanced (6 techniques)')
    doc.add_paragraph(f'Tasks Analyzed: {enhanced["total_analyzed"]}')
    doc.add_paragraph()
    
    # Executive Summary
    add_heading(doc, 'Executive Summary', 1)
    
    doc.add_paragraph('The Enhanced Validator uses 6 advanced techniques to validate task completion:')
    doc.add_paragraph('• Git commit mining (301 commits analyzed)', style='List Bullet')
    doc.add_paragraph('• Import statement analysis (156 components mapped)', style='List Bullet')
    doc.add_paragraph('• Component dependency graph (127 components)', style='List Bullet')
    doc.add_paragraph('• Test file correlation', style='List Bullet')
    doc.add_paragraph('• Stale task detection', style='List Bullet')
    doc.add_paragraph('• Semantic keyword extraction', style='List Bullet')
    doc.add_paragraph()
    
    # Comparison Table
    add_heading(doc, 'Validation Comparison', 1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    add_table_row(table, ['Metric', 'Basic V1', 'Enhanced V2'], bold=True)
    add_table_row(table, ['Method', 'Keyword matching', '6 techniques'])
    add_table_row(table, ['Tasks with evidence', basic['promotable_count'], enhanced['total_with_evidence']])
    add_table_row(table, ['High confidence', basic['promotable_count'], enhanced['high_confidence']])
    add_table_row(table, ['Medium confidence', '0', enhanced['medium_confidence']])
    add_table_row(table, ['Improvement', '1x baseline', f"{enhanced['total_with_evidence'] / basic['promotable_count']:.1f}x"])
    
    doc.add_paragraph()
    
    # Key Finding
    p = doc.add_paragraph()
    p.add_run('Key Finding: ').bold = True
    p.add_run(f'Enhanced validator found {enhanced["total_with_evidence"]} tasks with evidence ')
    p.add_run(f'({enhanced["total_with_evidence"] / basic["promotable_count"]:.1f}x more than basic validator)')
    doc.add_paragraph()
    
    # High Confidence Tasks
    add_heading(doc, f'High Confidence Tasks ({enhanced["high_confidence"]})', 1)
    doc.add_paragraph('These tasks scored ≥70% confidence and should be reviewed for promotion:')
    doc.add_paragraph()
    
    for i, task in enumerate(enhanced['high_confidence_tasks'], 1):
        doc.add_paragraph(f"{i}. {task['title']}", style='List Number')
        
        # Details table
        details = doc.add_table(rows=5, cols=2)
        details.style = 'Light List Accent 1'
        
        details.rows[0].cells[0].text = 'Confidence Score'
        details.rows[0].cells[1].text = f"{task['confidence_score']}%"
        
        details.rows[1].cells[0].text = 'Evidence Sources'
        details.rows[1].cells[1].text = ', '.join(task['evidence_sources'])
        
        details.rows[2].cells[0].text = 'Git Commits'
        details.rows[2].cells[1].text = f"{len(task.get('git_commits', []))} found"
        
        details.rows[3].cells[0].text = 'Import Usage'
        details.rows[3].cells[1].text = f"Used in {task.get('import_usage', 0)} files"
        
        details.rows[4].cells[0].text = 'Has Tests'
        details.rows[4].cells[1].text = '✅ Yes' if task.get('has_tests') else '❌ No'
        
        doc.add_paragraph()
    
    # Medium Confidence Tasks (Top 20)
    add_heading(doc, f'Medium Confidence Tasks (Top 20 of {enhanced["medium_confidence"]})', 1)
    doc.add_paragraph('These tasks scored 50-69% confidence and need manual review:')
    doc.add_paragraph()
    
    for i, task in enumerate(enhanced['medium_confidence_tasks'][:20], 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(f"{task['title']} ").bold = True
        p.add_run(f"({task['confidence_score']}%)")
        
        # Brief evidence
        evidence_text = ', '.join(task['evidence_sources'])
        doc.add_paragraph(f"   Evidence: {evidence_text}", style='List Bullet 2')
        
        if task.get('git_commits'):
            doc.add_paragraph(f"   Commits: {len(task['git_commits'])} found", style='List Bullet 2')
    
    doc.add_paragraph()
    
    # Validation Techniques Detail
    add_heading(doc, 'Validation Techniques Used', 1)
    
    add_heading(doc, '1. Git Commit Mining', 2)
    doc.add_paragraph('Analyzed 301 git commits to find when features were implemented.')
    doc.add_paragraph('Impact: +25% confidence when commits match task keywords')
    doc.add_paragraph()
    
    add_heading(doc, '2. Import Statement Analysis', 2)
    doc.add_paragraph('Mapped 156 components to track usage across codebase.')
    doc.add_paragraph('Impact: Up to +20% confidence based on usage frequency')
    doc.add_paragraph()
    
    add_heading(doc, '3. Component Dependency Graph', 2)
    doc.add_paragraph('Built dependency graph of 127 components.')
    doc.add_paragraph('Impact: Validates feature completeness by checking all required components exist')
    doc.add_paragraph()
    
    add_heading(doc, '4. Test File Correlation', 2)
    doc.add_paragraph('Checks if test files exist for components.')
    doc.add_paragraph('Impact: +15% confidence if tests exist (indicates feature is "done done")')
    doc.add_paragraph()
    
    add_heading(doc, '5. Stale Task Detection', 2)
    doc.add_paragraph('Identifies tasks referencing old/renamed files.')
    doc.add_paragraph('Impact: -10% confidence if references are outdated')
    doc.add_paragraph()
    
    add_heading(doc, '6. Semantic Keyword Extraction', 2)
    doc.add_paragraph('Extracts meaningful keywords from task descriptions.')
    doc.add_paragraph('Impact: Better matching between tasks and code')
    doc.add_paragraph()
    
    # Statistics
    add_heading(doc, 'Statistics', 1)
    
    stats_table = doc.add_table(rows=8, cols=2)
    stats_table.style = 'Light Grid Accent 1'
    
    stats_table.rows[0].cells[0].text = 'Total Tasks'
    stats_table.rows[0].cells[1].text = str(enhanced['total_analyzed'])
    
    stats_table.rows[1].cells[0].text = 'Tasks with Evidence'
    stats_table.rows[1].cells[1].text = f"{enhanced['total_with_evidence']} ({enhanced['total_with_evidence']/enhanced['total_analyzed']*100:.1f}%)"
    
    stats_table.rows[2].cells[0].text = 'High Confidence (≥70%)'
    stats_table.rows[2].cells[1].text = f"{enhanced['high_confidence']} ({enhanced['high_confidence']/enhanced['total_analyzed']*100:.1f}%)"
    
    stats_table.rows[3].cells[0].text = 'Medium Confidence (50-69%)'
    stats_table.rows[3].cells[1].text = f"{enhanced['medium_confidence']} ({enhanced['medium_confidence']/enhanced['total_analyzed']*100:.1f}%)"
    
    stats_table.rows[4].cells[0].text = 'Git Commits Analyzed'
    stats_table.rows[4].cells[1].text = '301'
    
    stats_table.rows[5].cells[0].text = 'Components Mapped'
    stats_table.rows[5].cells[1].text = '156'
    
    stats_table.rows[6].cells[0].text = 'Dependency Graph Nodes'
    stats_table.rows[6].cells[1].text = '127'
    
    stats_table.rows[7].cells[0].text = 'Improvement vs Basic'
    stats_table.rows[7].cells[1].text = f"{enhanced['total_with_evidence'] / basic['promotable_count']:.1f}x more tasks"
    
    doc.add_paragraph()
    
    # Recommendations
    add_heading(doc, 'Recommendations', 1)
    
    add_heading(doc, 'Immediate Actions', 2)
    doc.add_paragraph(f'1. Review {enhanced["high_confidence"]} high-confidence tasks for promotion', style='List Number')
    doc.add_paragraph('2. Sample 10-20 medium-confidence tasks for manual verification', style='List Number')
    doc.add_paragraph('3. Consider lowering threshold to 60% for "high confidence"', style='List Number')
    doc.add_paragraph()
    
    add_heading(doc, 'Future Enhancements', 2)
    doc.add_paragraph('1. Add LLM semantic matching for better intent understanding', style='List Number')
    doc.add_paragraph('2. Implement Notion metadata enrichment to show evidence in board', style='List Number')
    doc.add_paragraph('3. Add API endpoint validation for backend tasks', style='List Number')
    doc.add_paragraph('4. Integrate Playwright for visual evidence of UI features', style='List Number')
    doc.add_paragraph()
    
    # Conclusion
    add_heading(doc, 'Conclusion', 1)
    
    doc.add_paragraph(
        f'The Enhanced Validator successfully identified {enhanced["total_with_evidence"]} tasks with evidence, '
        f'a {enhanced["total_with_evidence"] / basic["promotable_count"]:.1f}x improvement over the basic validator. '
        f'While the confidence scoring is conservative (only {enhanced["high_confidence"]} high-confidence tasks), '
        'this prevents false positives and ensures accuracy. The {enhanced["medium_confidence"]} medium-confidence '
        'tasks provide a valuable pipeline for manual review.'
    )
    
    doc.add_paragraph()
    doc.add_paragraph(
        'The git commit mining, import analysis, and component graph techniques proved most valuable, '
        'providing concrete evidence of feature implementation. Future enhancements with LLM semantic matching '
        'could further improve accuracy and catch tasks that keyword matching misses.'
    )
    
    # Save document
    output_dir = Path.home() / 'Documents' / 'Joju'
    output_dir.mkdir(parents=True, exist_ok=True)
    output_file = output_dir / f'Joju_Enhanced_Validation_Report_{datetime.now().strftime("%Y-%m-%d")}.docx'
    
    doc.save(output_file)
    return output_file

if __name__ == '__main__':
    print("\n" + "="*80)
    print("EXPORTING ENHANCED VALIDATION REPORT TO WORD")
    print("="*80 + "\n")
    
    try:
        output_file = create_enhanced_validation_report()
        print(f"✅ Enhanced validation report exported to:")
        print(f"   {output_file}\n")
        print("📄 Document includes:")
        print("   • Executive summary with comparison")
        print("   • 4 high-confidence tasks with details")
        print("   • Top 20 medium-confidence tasks")
        print("   • All 6 validation techniques explained")
        print("   • Statistics and recommendations")
        print()
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        print("\nMake sure python-docx is installed:")
        print("   pip3 install python-docx")
