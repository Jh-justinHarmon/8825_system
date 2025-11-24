#!/usr/bin/env python3
"""
Template-Based Word Document Generator V2
Creates documents by copying template header and page setup, then adding new content
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from pathlib import Path
import json
import sys
from typing import Dict, Any, Optional


def copy_header_footer(source_doc, target_doc):
    """Copy header and footer from source to target document"""
    try:
        # Copy headers
        for section_idx, source_section in enumerate(source_doc.sections):
            if section_idx < len(target_doc.sections):
                target_section = target_doc.sections[section_idx]
                
                # Copy header
                if source_section.header:
                    target_header = target_section.header
                    # Clear existing header
                    for element in list(target_header._element):
                        target_header._element.remove(element)
                    # Copy header content
                    for element in source_section.header._element:
                        target_header._element.append(element)
                
                # Copy footer  
                if source_section.footer:
                    target_footer = target_section.footer
                    # Clear existing footer
                    for element in list(target_footer._element):
                        target_footer._element.remove(element)
                    # Copy footer content
                    for element in source_section.footer._element:
                        target_footer._element.append(element)
    except Exception as e:
        print(f"Warning: Could not copy headers/footers: {e}")


def copy_page_setup(source_doc, target_doc):
    """Copy page setup (margins, size) from source to target"""
    for source_section, target_section in zip(source_doc.sections, target_doc.sections):
        target_section.page_width = source_section.page_width
        target_section.page_height = source_section.page_height
        target_section.left_margin = source_section.left_margin
        target_section.right_margin = source_section.right_margin
        target_section.top_margin = source_section.top_margin
        target_section.bottom_margin = source_section.bottom_margin


def generate_from_dict(template_path: str, content: Dict[str, Any], output_path: str) -> Path:
    """Generate Word document from template and content dict"""
    
    # Load template to get styles and setup
    template_doc = Document(template_path)
    
    # Create new blank document
    doc = Document()
    
    # Copy page setup from template
    copy_page_setup(template_doc, doc)
    
    # Copy header/footer from template
    copy_header_footer(template_doc, doc)
    
    # Add title
    if 'title' in content:
        title = doc.add_heading(content['title'], 0)
        if 'title_alignment' in content:
            if content['title_alignment'] == 'center':
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif content['title_alignment'] == 'right':
                title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add subtitle
    if 'subtitle' in content:
        subtitle = doc.add_paragraph(content['subtitle'])
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in subtitle.runs:
            run.italic = True
    
    doc.add_paragraph()  # Spacing
    
    # Add sections
    for section in content.get('sections', []):
        # Section heading
        if 'heading' in section:
            level = section.get('level', 1)
            doc.add_heading(section['heading'], level)
        
        # Paragraphs
        for para_text in section.get('paragraphs', []):
            doc.add_paragraph(para_text)
        
        # Bullet points
        for bullet in section.get('bullets', []):
            p = doc.add_paragraph()
            p.add_run('• ' + bullet)
        
        # Numbered lists
        for i, item in enumerate(section.get('numbered', []), 1):
            doc.add_paragraph(f"{i}. {item}")
        
        # Tables
        if 'table' in section:
            table_data = section['table']
            rows = table_data.get('rows', [])
            
            if rows:
                cols = len(rows[0])
                table = doc.add_table(rows=len(rows), cols=cols)
                
                # Populate table
                for i, row_data in enumerate(rows):
                    row_cells = table.rows[i].cells
                    for j, cell_data in enumerate(row_data):
                        row_cells[j].text = str(cell_data)
                        
                        # Bold header row
                        if i == 0 and table_data.get('header', True):
                            for paragraph in row_cells[j].paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
            
            doc.add_paragraph()  # Spacing after table
        
        doc.add_paragraph()  # Spacing between sections
    
    # Add footer if provided
    if 'footer' in content:
        doc.add_paragraph()
        footer = doc.add_paragraph(content['footer'])
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.italic = True
    
    # Save document
    output_path = Path(output_path)
    doc.save(output_path)
    
    return output_path


def main():
    """Command line interface"""
    if len(sys.argv) < 4:
        print("Template-Based Word Document Generator V2")
        print("\nUsage:")
        print("  python3 template_word_generator_v2.py <template.docx> <content.json> <output.docx>")
        sys.exit(1)
    
    template_path = sys.argv[1]
    content_path = sys.argv[2]
    output_path = sys.argv[3]
    
    try:
        with open(content_path, 'r') as f:
            content = json.load(f)
        
        output = generate_from_dict(template_path, content, output_path)
        
        print(f"✅ Generated: {output}")
        print(f"   Template: {template_path}")
        print(f"   Content: {content_path}")
    
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
