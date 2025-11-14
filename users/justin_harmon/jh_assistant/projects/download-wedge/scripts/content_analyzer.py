#!/usr/bin/env python3
"""
Content Analyzer
Analyzes file content using text extraction, PDF parsing, and OCR
"""

import os
import re
from pathlib import Path

def analyze_content(file_path, metadata):
    """
    Analyze file content based on file type
    
    Args:
        file_path: Path to file
        metadata: Metadata dict from metadata_extractor
    
    Returns:
        dict: Content analysis results
    """
    category = metadata.get('category', 'other')
    extension = metadata.get('extension', '')
    
    content_data = {
        "text_sample": "",
        "keywords": [],
        "entities": [],
        "has_content": False
    }
    
    try:
        # Text files
        if extension in ['.txt', '.md', '.json', '.py', '.sh', '.html', '.css', '.js']:
            content_data = analyze_text_file(file_path)
        
        # PDF files
        elif extension == '.pdf':
            content_data = analyze_pdf(file_path)
        
        # DOCX files
        elif extension == '.docx':
            content_data = analyze_docx(file_path)
        
        # Images (OCR)
        elif category == 'image':
            content_data = analyze_image(file_path)
        
    except Exception as e:
        content_data["error"] = str(e)
    
    return content_data

def analyze_text_file(file_path):
    """Analyze plain text file"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            # Read first 5000 characters
            content = f.read(5000)
        
        return {
            "text_sample": content[:500],  # First 500 chars
            "keywords": extract_keywords(content),
            "entities": extract_entities(content),
            "has_content": len(content.strip()) > 0
        }
    except:
        return {"has_content": False}

def analyze_pdf(file_path):
    """Analyze PDF file"""
    try:
        import PyPDF2
        
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            
            # Extract text from first page
            if len(reader.pages) > 0:
                text = reader.pages[0].extract_text()
                
                return {
                    "text_sample": text[:500],
                    "keywords": extract_keywords(text),
                    "entities": extract_entities(text),
                    "has_content": len(text.strip()) > 0,
                    "page_count": len(reader.pages)
                }
    except:
        pass
    
    return {"has_content": False, "note": "PDF extraction not available"}

def analyze_docx(file_path):
    """Analyze DOCX file"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        
        # Extract text from first few paragraphs
        text = "\n".join([p.text for p in doc.paragraphs[:10]])
        
        return {
            "text_sample": text[:500],
            "keywords": extract_keywords(text),
            "entities": extract_entities(text),
            "has_content": len(text.strip()) > 0,
            "paragraph_count": len(doc.paragraphs)
        }
    except:
        pass
    
    return {"has_content": False, "note": "DOCX extraction not available"}

def analyze_image(file_path):
    """Analyze image using OCR"""
    try:
        import pytesseract
        from PIL import Image
        
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        
        return {
            "text_sample": text[:500],
            "keywords": extract_keywords(text),
            "entities": extract_entities(text),
            "has_content": len(text.strip()) > 0,
            "ocr_used": True
        }
    except:
        pass
    
    return {"has_content": False, "note": "OCR not available"}

def extract_keywords(text):
    """
    Extract relevant keywords from text
    
    Returns:
        list: Detected keywords
    """
    keywords = []
    
    text_lower = text.lower()
    
    # Project keywords
    project_keywords = [
        'tgif', 'meeting', 'summary', 'action items',
        'ral', 'design', 'mockup', 'prototype',
        'hcss', 'hammer consulting', 'consulting',
        'joju', 'forge', 'trustybit', '76',
        'protocol', 'workflow', 'agent', '8825',
        'personal', 'justin', 'script'
    ]
    
    for keyword in project_keywords:
        if keyword in text_lower:
            keywords.append(keyword)
    
    return list(set(keywords))  # Remove duplicates

def extract_entities(text):
    """
    Extract entities (names, companies, etc.) from text
    
    Returns:
        list: Detected entities
    """
    entities = []
    
    # Simple entity patterns
    entity_patterns = [
        'Justin Harmon',
        'Hammer Consulting',
        'HCSS',
        'RAL',
        'Joju',
        'Forge',
        'TrustyBit'
    ]
    
    for entity in entity_patterns:
        if entity.lower() in text.lower():
            entities.append(entity)
    
    return entities

if __name__ == "__main__":
    # Test
    import sys
    import json
    from metadata_extractor import extract_metadata
    
    if len(sys.argv) > 1:
        metadata = extract_metadata(sys.argv[1])
        content = analyze_content(sys.argv[1], metadata)
        print(json.dumps(content, indent=2))
