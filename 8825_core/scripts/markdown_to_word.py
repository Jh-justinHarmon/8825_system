#!/usr/bin/env python3
"""
Convert Markdown files to Word documents
Generic utility for exporting markdown to professional Word format
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
import re
import sys
from pathlib import Path

def export_markdown_to_word(md_file, docx_file=None):
    """
    Convert markdown to Word document
    
    Args:
        md_file: Path to markdown file
        docx_file: Output path (optional, defaults to same name with .docx)
    
    Returns:
        Path to created Word document
    """
    
    md_path = Path(md_file)
    
    if not md_path.exists():
        raise FileNotFoundError(f"Markdown file not found: {md_file}")
    
    # Default output path
    if docx_file is None:
        docx_file = md_path.with_suffix('.docx')
    else:
        docx_file = Path(docx_file)
    
    # Read markdown
    with open(md_path, 'r') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Process content
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines at start
        if not line:
            i += 1
            continue
        
        # H1 headers
        if line.startswith('# '):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        # H2 headers
        elif line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_heading(text, level=2)
            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        # H3 headers
        elif line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_heading(text, level=3)
            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        # Horizontal rules
        elif line.startswith('---'):
            doc.add_paragraph()
        
        # Code blocks
        elif line.startswith('```'):
            # Collect code block
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            # Add code block
            if code_lines:
                p = doc.add_paragraph()
                p.style = 'No Spacing'
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
        
        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            # Remove markdown bold
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'__(.*?)__', r'\1', text)
            doc.add_paragraph(text, style='List Bullet')
        
        # Regular paragraphs
        elif line:
            # Remove markdown formatting
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            text = re.sub(r'__(.*?)__', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(6)
        
        i += 1
    
    # Save
    doc.save(docx_file)
    return docx_file

def main():
    """Command line interface"""
    if len(sys.argv) < 2:
        print("Usage: python3 markdown_to_word.py <markdown_file> [output_file]")
        print("\nExample:")
        print("  python3 markdown_to_word.py README.md")
        print("  python3 markdown_to_word.py README.md output.docx")
        sys.exit(1)
    
    md_file = sys.argv[1]
    docx_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    try:
        output_path = export_markdown_to_word(md_file, docx_file)
        print(f'✅ Exported to: {output_path}')
    except Exception as e:
        print(f'❌ Error: {e}')
        sys.exit(1)

if __name__ == '__main__':
    main()
