#!/usr/bin/env python3
"""
Analyze user testing documents and extract key insights
"""

import os
from pathlib import Path
from docx import Document
from datetime import datetime
from collections import defaultdict

def extract_text_from_docx(filepath):
    """Extract all text from a Word document"""
    doc = Document(filepath)
    text = []
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text)
    return '\n'.join(text)

def analyze_session(filepath):
    """Analyze a single user testing session"""
    filename = Path(filepath).name
    
    # Extract participant name from filename
    participant = filename.split(' User')[0].strip()
    
    # Extract date
    date_str = None
    if '2025_08_' in filename:
        date_part = filename.split('2025_08_')[1].split(' ')[0]
        date_str = f"2025-08-{date_part}"
    
    # Extract text
    text = extract_text_from_docx(filepath)
    
    # Simple keyword extraction
    insights = {
        'participant': participant,
        'date': date_str,
        'filepath': str(filepath),
        'word_count': len(text.split()),
        'key_themes': [],
        'pain_points': [],
        'positive_feedback': [],
        'suggestions': []
    }
    
    # Look for common patterns (simple keyword matching)
    lines = text.lower().split('\n')
    
    for line in lines:
        # Pain points
        if any(word in line for word in ['difficult', 'confusing', 'problem', 'issue', 'frustrated', 'hard to']):
            insights['pain_points'].append(line.strip())
        
        # Positive feedback
        if any(word in line for word in ['easy', 'like', 'helpful', 'good', 'great', 'love', 'intuitive']):
            insights['positive_feedback'].append(line.strip())
        
        # Suggestions
        if any(word in line for word in ['suggest', 'recommend', 'would be better', 'should', 'could']):
            insights['suggestions'].append(line.strip())
    
    # Limit to top items
    insights['pain_points'] = insights['pain_points'][:5]
    insights['positive_feedback'] = insights['positive_feedback'][:5]
    insights['suggestions'] = insights['suggestions'][:5]
    
    return insights

def create_summary_document(all_insights, output_path):
    """Create a comprehensive summary document"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Joju Intake Design Sprint', 0)
    title.alignment = 1  # Center
    
    subtitle = doc.add_paragraph()
    subtitle.add_run('User Testing Summary Report\n').bold = True
    subtitle.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}\n').italic = True
    subtitle.add_run(f'Sessions Analyzed: {len(all_insights)}').italic = True
    subtitle.alignment = 1
    
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading('📊 Executive Summary', 1)
    
    total_words = sum(s['word_count'] for s in all_insights)
    total_pain_points = sum(len(s['pain_points']) for s in all_insights)
    total_positive = sum(len(s['positive_feedback']) for s in all_insights)
    total_suggestions = sum(len(s['suggestions']) for s in all_insights)
    
    summary = doc.add_paragraph()
    summary.add_run(f'• Total Sessions: {len(all_insights)}\n')
    summary.add_run(f'• Total Content: {total_words:,} words\n')
    summary.add_run(f'• Pain Points Identified: {total_pain_points}\n')
    summary.add_run(f'• Positive Feedback Items: {total_positive}\n')
    summary.add_run(f'• Suggestions Captured: {total_suggestions}\n')
    
    doc.add_paragraph()
    
    # Sessions Overview
    doc.add_heading('👥 Testing Sessions', 1)
    
    for insight in sorted(all_insights, key=lambda x: x['date'] or ''):
        doc.add_heading(f"{insight['participant']}", 2)
        info = doc.add_paragraph()
        info.add_run(f"Date: {insight['date']}\n")
        info.add_run(f"Content: {insight['word_count']:,} words\n")
    
    doc.add_paragraph()
    
    # Aggregated Pain Points
    doc.add_heading('⚠️ Key Pain Points', 1)
    
    all_pain_points = []
    for insight in all_insights:
        for pain in insight['pain_points']:
            all_pain_points.append(f"[{insight['participant']}] {pain}")
    
    if all_pain_points:
        for i, pain in enumerate(all_pain_points[:15], 1):
            doc.add_paragraph(f"{i}. {pain}", style='List Number')
    else:
        doc.add_paragraph("No specific pain points extracted.")
    
    doc.add_paragraph()
    
    # Positive Feedback
    doc.add_heading('✅ Positive Feedback', 1)
    
    all_positive = []
    for insight in all_insights:
        for pos in insight['positive_feedback']:
            all_positive.append(f"[{insight['participant']}] {pos}")
    
    if all_positive:
        for i, pos in enumerate(all_positive[:15], 1):
            doc.add_paragraph(f"{i}. {pos}", style='List Number')
    else:
        doc.add_paragraph("No specific positive feedback extracted.")
    
    doc.add_paragraph()
    
    # Suggestions
    doc.add_heading('💡 Suggestions & Recommendations', 1)
    
    all_suggestions = []
    for insight in all_insights:
        for sug in insight['suggestions']:
            all_suggestions.append(f"[{insight['participant']}] {sug}")
    
    if all_suggestions:
        for i, sug in enumerate(all_suggestions[:15], 1):
            doc.add_paragraph(f"{i}. {sug}", style='List Number')
    else:
        doc.add_paragraph("No specific suggestions extracted.")
    
    doc.add_paragraph()
    
    # Individual Session Details
    doc.add_page_break()
    doc.add_heading('📋 Individual Session Details', 1)
    
    for insight in sorted(all_insights, key=lambda x: x['date'] or ''):
        doc.add_heading(f"{insight['participant']} - {insight['date']}", 2)
        
        if insight['pain_points']:
            doc.add_heading('Pain Points:', 3)
            for pain in insight['pain_points']:
                doc.add_paragraph(pain, style='List Bullet')
        
        if insight['positive_feedback']:
            doc.add_heading('Positive Feedback:', 3)
            for pos in insight['positive_feedback']:
                doc.add_paragraph(pos, style='List Bullet')
        
        if insight['suggestions']:
            doc.add_heading('Suggestions:', 3)
            for sug in insight['suggestions']:
                doc.add_paragraph(sug, style='List Bullet')
        
        doc.add_paragraph()
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run('Generated by 8825 User Testing Analysis Pipeline').italic = True
    footer.alignment = 1
    
    # Save
    doc.save(output_path)
    return output_path

def main():
    downloads_dir = Path.home() / 'Downloads'
    
    print(f"\n{'='*80}")
    print(f"🔍 Analyzing Joju User Testing Sessions")
    print(f"{'='*80}\n")
    
    # Find all user testing documents
    patterns = [
        '*User Testing*.docx',
        '*User Test*.docx'
    ]
    
    files = []
    for pattern in patterns:
        files.extend(downloads_dir.glob(pattern))
    
    if not files:
        print("❌ No user testing documents found in Downloads")
        return
    
    print(f"Found {len(files)} document(s):\n")
    
    all_insights = []
    
    for filepath in sorted(files):
        print(f"Analyzing: {filepath.name}")
        insights = analyze_session(filepath)
        all_insights.append(insights)
        
        print(f"  • Participant: {insights['participant']}")
        print(f"  • Date: {insights['date']}")
        print(f"  • Words: {insights['word_count']:,}")
        print(f"  • Pain points: {len(insights['pain_points'])}")
        print(f"  • Positive feedback: {len(insights['positive_feedback'])}")
        print(f"  • Suggestions: {len(insights['suggestions'])}")
        print()
    
    # Create summary document
    print(f"{'─'*80}")
    print("Creating summary document...")
    
    output_path = downloads_dir / f"Joju_User_Testing_Summary_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    create_summary_document(all_insights, output_path)
    
    print(f"✅ Summary created: {output_path.name}")
    print(f"\n{'='*80}")
    print(f"Analysis Complete!")
    print(f"{'='*80}\n")
    
    print(f"Summary Statistics:")
    print(f"  • Total sessions: {len(all_insights)}")
    print(f"  • Total words: {sum(s['word_count'] for s in all_insights):,}")
    print(f"  • Total pain points: {sum(len(s['pain_points']) for s in all_insights)}")
    print(f"  • Total positive feedback: {sum(len(s['positive_feedback']) for s in all_insights)}")
    print(f"  • Total suggestions: {sum(len(s['suggestions']) for s in all_insights)}")
    print(f"\nOutput: {output_path}")
    print()

if __name__ == "__main__":
    main()
