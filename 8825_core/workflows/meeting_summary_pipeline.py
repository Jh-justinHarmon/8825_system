#!/usr/bin/env python3
"""
HCSS Meeting Summary Pipeline
Processes meeting notes/transcripts and generates Word documents
"""

import json
import os
import sys
from pathlib import Path
from datetime import datetime, timedelta
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, asdict
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add file router to path
sys.path.insert(0, str(Path(__file__).parent.parent / "config"))
from file_router import get_destination


@dataclass
class MeetingData:
    """Meeting data structure"""
    title: str
    date: str
    attendees: List[str]
    topics: List[str]
    decisions: List[Dict[str, Any]]
    action_items: List[Dict[str, Any]]
    risks: List[Dict[str, Any]]
    blockers: List[Dict[str, Any]]
    next_steps: List[str]
    notes: Optional[str] = None
    source_file: Optional[str] = None


class MeetingSummaryPipeline:
    """Pipeline for processing meeting notes and generating Word docs"""
    
    def __init__(self, 
                 input_dir: str = None,
                 output_dir: str = None,
                 focus: str = "hcss"):
        """
        Initialize pipeline
        
        Args:
            input_dir: Directory to scan for meeting notes
            output_dir: Directory for output Word docs
            focus: Focus area (hcss, joju, etc.)
        """
        if input_dir is None:
            input_dir = get_destination("HCSS")
        
        if output_dir is None:
            output_dir = get_destination("HCSS") / "Summaries"
        
        self.input_dir = Path(input_dir)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.focus = focus
    
    def scan_for_meetings(self, start_date: datetime, end_date: datetime) -> List[Path]:
        """
        Scan input directory for meeting files in date range
        
        Args:
            start_date: Start of date range
            end_date: End of date range
            
        Returns:
            List of meeting file paths
        """
        meeting_files = []
        
        # Scan for JSON files
        for json_file in self.input_dir.rglob("*.json"):
            # Skip meta files
            if ".meta." in json_file.name:
                continue
            
            try:
                with open(json_file, 'r') as f:
                    data = json.load(f)
                    
                    # Check if it's a meeting note
                    if data.get("content_type") == "note" or "meeting" in json_file.name.lower():
                        # Extract date
                        meeting_date = self._extract_date(data, json_file)
                        
                        if meeting_date and start_date <= meeting_date <= end_date:
                            meeting_files.append(json_file)
            except Exception as e:
                print(f"Error reading {json_file}: {e}")
        
        return sorted(meeting_files)
    
    def _extract_date(self, data: Dict, filepath: Path) -> Optional[datetime]:
        """Extract date from meeting data or filename"""
        # Try content date
        if "content" in data and "date" in data["content"]:
            try:
                return datetime.fromisoformat(data["content"]["date"])
            except:
                pass
        
        # Try metadata timestamp
        if "metadata" in data and "timestamp" in data["metadata"]:
            try:
                ts = data["metadata"]["timestamp"]
                return datetime.fromisoformat(ts.replace('Z', '+00:00'))
            except:
                pass
        
        # Try filename
        filename = filepath.stem
        if filename.startswith("2025"):
            try:
                date_str = filename[:8]  # YYYYMMDD
                return datetime.strptime(date_str, "%Y%m%d")
            except:
                pass
        
        return None
    
    def parse_meeting_file(self, filepath: Path) -> MeetingData:
        """
        Parse meeting file into structured data
        
        Args:
            filepath: Path to meeting file
            
        Returns:
            MeetingData object
        """
        with open(filepath, 'r') as f:
            data = json.load(f)
        
        content = data.get("content", {})
        
        # Extract basic info
        title = content.get("title", "HCSS Meeting")
        date = content.get("date", datetime.now().strftime("%Y-%m-%d"))
        attendees = content.get("attendees", [])
        topics = content.get("topics", [])
        
        # Extract structured items
        decisions = []
        if "decisions" in content:
            for dec in content["decisions"]:
                if isinstance(dec, str):
                    decisions.append({"text": dec, "impact": "medium"})
                else:
                    decisions.append(dec)
        
        action_items = []
        if "action_items" in content:
            for action in content["action_items"]:
                if isinstance(action, str):
                    action_items.append({
                        "what": action,
                        "who": "TBD",
                        "due": None,
                        "priority": "medium"
                    })
                else:
                    action_items.append(action)
        
        risks = content.get("risks", [])
        blockers = content.get("blockers", [])
        next_steps = content.get("next_steps", [])
        notes = content.get("notes", "")
        
        return MeetingData(
            title=title,
            date=date,
            attendees=attendees,
            topics=topics,
            decisions=decisions,
            action_items=action_items,
            risks=risks,
            blockers=blockers,
            next_steps=next_steps,
            notes=notes,
            source_file=str(filepath)
        )
    
    def generate_word_doc(self, meeting: MeetingData) -> Path:
        """
        Generate Word document for meeting
        
        Args:
            meeting: MeetingData object
            
        Returns:
            Path to generated Word doc
        """
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Title
        title = doc.add_heading(meeting.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Meeting info
        info_para = doc.add_paragraph()
        info_para.add_run(f"Date: ").bold = True
        info_para.add_run(f"{meeting.date}\n")
        info_para.add_run(f"Attendees: ").bold = True
        info_para.add_run(f"{', '.join(meeting.attendees)}\n")
        
        doc.add_paragraph()  # Spacing
        
        # Topics Discussed
        if meeting.topics:
            doc.add_heading("Topics Discussed", 1)
            for topic in meeting.topics:
                doc.add_paragraph(topic, style='List Bullet')
            doc.add_paragraph()
        
        # Decisions
        if meeting.decisions:
            doc.add_heading("🎯 Decisions", 1)
            for i, decision in enumerate(meeting.decisions, 1):
                if isinstance(decision, dict):
                    text = decision.get("text", str(decision))
                    impact = decision.get("impact", "medium")
                    para = doc.add_paragraph(f"{i}. {text}", style='List Number')
                    if impact in ["high", "critical"]:
                        para.add_run(f" [Impact: {impact.upper()}]").bold = True
                else:
                    doc.add_paragraph(f"{i}. {decision}", style='List Number')
            doc.add_paragraph()
        
        # Action Items
        if meeting.action_items:
            doc.add_heading("✅ Action Items", 1)
            
            # Create table
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Action'
            header_cells[1].text = 'Owner'
            header_cells[2].text = 'Due Date'
            header_cells[3].text = 'Priority'
            
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Add action items
            for action in meeting.action_items:
                row_cells = table.add_row().cells
                if isinstance(action, dict):
                    row_cells[0].text = action.get("what", str(action))
                    row_cells[1].text = action.get("who", "TBD")
                    row_cells[2].text = action.get("due", "TBD") or "TBD"
                    row_cells[3].text = action.get("priority", "medium").upper()
                else:
                    row_cells[0].text = str(action)
                    row_cells[1].text = "TBD"
                    row_cells[2].text = "TBD"
                    row_cells[3].text = "MEDIUM"
            
            doc.add_paragraph()
        
        # Risks
        if meeting.risks:
            doc.add_heading("⚠️ Risks", 1)
            for i, risk in enumerate(meeting.risks, 1):
                if isinstance(risk, dict):
                    text = risk.get("text", str(risk))
                    severity = risk.get("severity", "medium")
                    para = doc.add_paragraph(f"{i}. {text}", style='List Number')
                    para.add_run(f" [Severity: {severity.upper()}]").bold = True
                else:
                    doc.add_paragraph(f"{i}. {risk}", style='List Number')
            doc.add_paragraph()
        
        # Blockers
        if meeting.blockers:
            doc.add_heading("🚫 Blockers", 1)
            for i, blocker in enumerate(meeting.blockers, 1):
                if isinstance(blocker, dict):
                    text = blocker.get("text", str(blocker))
                    impact = blocker.get("impact", "")
                    para = doc.add_paragraph(f"{i}. {text}", style='List Number')
                    if impact:
                        para.add_run(f" [Impact: {impact}]").italic = True
                else:
                    doc.add_paragraph(f"{i}. {blocker}", style='List Number')
            doc.add_paragraph()
        
        # Next Steps
        if meeting.next_steps:
            doc.add_heading("📅 Next Steps", 1)
            for step in meeting.next_steps:
                doc.add_paragraph(step, style='List Bullet')
            doc.add_paragraph()
        
        # Notes
        if meeting.notes:
            doc.add_heading("📝 Additional Notes", 1)
            doc.add_paragraph(meeting.notes)
        
        # Footer
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.add_run("Generated by 8825 Meeting Summary Pipeline").italic = True
        footer_para.add_run(f"\nGenerated: {datetime.now().strftime('%Y-%m-%d %H:%M')}").italic = True
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save document
        filename = f"{meeting.date}_{meeting.title.replace(' ', '_')}.docx"
        filepath = self.output_dir / filename
        doc.save(filepath)
        
        return filepath
    
    def process_last_week(self) -> List[Path]:
        """
        Process all meetings from last week and generate Word docs
        
        Returns:
            List of generated Word doc paths
        """
        # Calculate last week's date range (Monday to Friday)
        today = datetime.now().date()
        
        # Get last Monday
        days_since_monday = (today.weekday() + 7) % 7  # 0 = Monday
        if days_since_monday == 0:
            # If today is Monday, go back to last week
            last_monday = today - timedelta(days=7)
        else:
            last_monday = today - timedelta(days=days_since_monday + 7)
        
        last_friday = last_monday + timedelta(days=4)
        
        print(f"\n{'='*60}")
        print(f"Processing meetings from {last_monday} to {last_friday}")
        print(f"{'='*60}\n")
        
        # Convert to datetime for comparison
        start_date = datetime.combine(last_monday, datetime.min.time())
        end_date = datetime.combine(last_friday, datetime.max.time())
        
        # Scan for meetings
        meeting_files = self.scan_for_meetings(start_date, end_date)
        
        if not meeting_files:
            print(f"⚠️  No meetings found for last week ({last_monday} to {last_friday})")
            print(f"   Searched in: {self.input_dir}")
            return []
        
        print(f"Found {len(meeting_files)} meeting(s):\n")
        
        # Process each meeting
        generated_docs = []
        for meeting_file in meeting_files:
            print(f"Processing: {meeting_file.name}")
            
            try:
                # Parse meeting
                meeting = self.parse_meeting_file(meeting_file)
                
                # Generate Word doc
                doc_path = self.generate_word_doc(meeting)
                generated_docs.append(doc_path)
                
                print(f"✅ Generated: {doc_path.name}\n")
            
            except Exception as e:
                print(f"❌ Error processing {meeting_file.name}: {e}\n")
        
        print(f"\n{'='*60}")
        print(f"Summary: Generated {len(generated_docs)} Word document(s)")
        print(f"Output directory: {self.output_dir}")
        print(f"{'='*60}\n")
        
        return generated_docs
    
    def generate_weekly_summary(self, meetings: List[MeetingData]) -> Path:
        """
        Generate weekly summary document
        
        Args:
            meetings: List of MeetingData objects
            
        Returns:
            Path to weekly summary Word doc
        """
        doc = Document()
        
        # Title
        title = doc.add_heading("HCSS Weekly Meeting Summary", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Week info
        if meetings:
            dates = [m.date for m in meetings]
            week_start = min(dates)
            week_end = max(dates)
            
            info_para = doc.add_paragraph()
            info_para.add_run(f"Week: ").bold = True
            info_para.add_run(f"{week_start} to {week_end}\n")
            info_para.add_run(f"Meetings: ").bold = True
            info_para.add_run(f"{len(meetings)}\n")
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Aggregate all items
        all_decisions = []
        all_actions = []
        all_risks = []
        all_blockers = []
        
        for meeting in meetings:
            all_decisions.extend(meeting.decisions)
            all_actions.extend(meeting.action_items)
            all_risks.extend(meeting.risks)
            all_blockers.extend(meeting.blockers)
        
        # Summary stats
        doc.add_heading("📊 Week at a Glance", 1)
        stats_para = doc.add_paragraph()
        stats_para.add_run(f"• Total Decisions: {len(all_decisions)}\n")
        stats_para.add_run(f"• Total Action Items: {len(all_actions)}\n")
        stats_para.add_run(f"• Total Risks: {len(all_risks)}\n")
        stats_para.add_run(f"• Total Blockers: {len(all_blockers)}\n")
        
        doc.add_paragraph()
        
        # Meeting summaries
        doc.add_heading("📅 Meetings This Week", 1)
        for meeting in meetings:
            doc.add_heading(f"{meeting.date} - {meeting.title}", 2)
            doc.add_paragraph(f"Attendees: {', '.join(meeting.attendees)}")
            if meeting.topics:
                doc.add_paragraph(f"Topics: {', '.join(meeting.topics)}")
            doc.add_paragraph()
        
        # All action items
        if all_actions:
            doc.add_heading("✅ All Action Items", 1)
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Action'
            header_cells[1].text = 'Owner'
            header_cells[2].text = 'Due Date'
            header_cells[3].text = 'Priority'
            
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            for action in all_actions:
                row_cells = table.add_row().cells
                if isinstance(action, dict):
                    row_cells[0].text = action.get("what", str(action))
                    row_cells[1].text = action.get("who", "TBD")
                    row_cells[2].text = action.get("due", "TBD") or "TBD"
                    row_cells[3].text = action.get("priority", "medium").upper()
                else:
                    row_cells[0].text = str(action)
                    row_cells[1].text = "TBD"
                    row_cells[2].text = "TBD"
                    row_cells[3].text = "MEDIUM"
        
        # Save
        filename = f"Weekly_Summary_{datetime.now().strftime('%Y%m%d')}.docx"
        filepath = self.output_dir / filename
        doc.save(filepath)
        
        return filepath


def main():
    """Main entry point"""
    pipeline = MeetingSummaryPipeline()
    
    # Process last week's meetings
    docs = pipeline.process_last_week()
    
    if docs:
        print(f"\n✅ Success! Generated {len(docs)} document(s)")
        print(f"\nDocuments saved to:")
        print(f"  {pipeline.output_dir}")
    else:
        print("\n⚠️  No meetings found to process")
        print("\nTo add meetings, place JSON files in:")
        print(f"  {pipeline.input_dir}")
        print("\nExpected format:")
        print("""
{
  "content_type": "note",
  "target_focus": "hcss",
  "content": {
    "title": "Meeting Title",
    "date": "2025-11-04",
    "attendees": ["Person 1", "Person 2"],
    "topics": ["Topic 1", "Topic 2"],
    "action_items": ["Action 1", "Action 2"],
    "decisions": ["Decision 1"],
    "risks": [],
    "blockers": [],
    "next_steps": []
  }
}
        """)


if __name__ == "__main__":
    main()
