#!/usr/bin/env python3
"""
Export all action items to a Word document
"""

from meeting_summary_pipeline import MeetingSummaryPipeline
from datetime import datetime, timedelta
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    """Extract all action items and export to Word"""
    pipeline = MeetingSummaryPipeline()
    
    # Get date range for past 2 weeks
    today = datetime.now().date()
    two_weeks_ago = today - timedelta(days=14)
    
    start_date = datetime.combine(two_weeks_ago, datetime.min.time())
    end_date = datetime.combine(today, datetime.max.time())
    
    # Scan for meetings
    meeting_files = pipeline.scan_for_meetings(start_date, end_date)
    
    if not meeting_files:
        print("No meetings found")
        return
    
    # Collect all action items
    all_actions = []
    
    for meeting_file in meeting_files:
        try:
            meeting = pipeline.parse_meeting_file(meeting_file)
            
            for action in meeting.action_items:
                if isinstance(action, dict):
                    all_actions.append({
                        'meeting': meeting.title,
                        'date': meeting.date,
                        'what': action.get('what', str(action)),
                        'who': action.get('who', 'TBD'),
                        'due': action.get('due', 'TBD'),
                        'priority': action.get('priority', 'medium').upper()
                    })
                else:
                    all_actions.append({
                        'meeting': meeting.title,
                        'date': meeting.date,
                        'what': str(action),
                        'who': 'TBD',
                        'due': 'TBD',
                        'priority': 'MEDIUM'
                    })
        except Exception as e:
            print(f"Error processing {meeting_file.name}: {e}")
    
    # Sort by priority and due date
    priority_order = {'CRITICAL': 0, 'HIGH': 1, 'MEDIUM': 2, 'LOW': 3}
    all_actions.sort(key=lambda x: (
        priority_order.get(x['priority'], 4),
        x['due'] if x['due'] and x['due'] != 'TBD' else '9999-99-99'
    ))
    
    # Create Word document
    doc = Document()
    
    # Title
    title = doc.add_heading('HCSS Complete Task List', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}\n').italic = True
    subtitle.add_run(f'Date Range: {two_weeks_ago} to {today}\n').italic = True
    subtitle.add_run(f'Total Tasks: {len(all_actions)}').bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Summary stats
    doc.add_heading('📊 Summary by Priority', 1)
    
    priority_counts = {}
    for action in all_actions:
        priority = action['priority']
        priority_counts[priority] = priority_counts.get(priority, 0) + 1
    
    stats_para = doc.add_paragraph()
    for priority in ['CRITICAL', 'HIGH', 'MEDIUM', 'LOW']:
        count = priority_counts.get(priority, 0)
        if count > 0:
            emoji = '🔴' if priority == 'CRITICAL' else '🟠' if priority == 'HIGH' else '🟡' if priority == 'MEDIUM' else '🟢'
            stats_para.add_run(f'{emoji} {priority}: {count} tasks\n')
    
    doc.add_paragraph()
    
    # All tasks by priority
    doc.add_heading('✅ All Action Items by Priority', 1)
    
    for priority in ['CRITICAL', 'HIGH', 'MEDIUM', 'LOW']:
        priority_tasks = [a for a in all_actions if a['priority'] == priority]
        
        if priority_tasks:
            # Priority heading
            emoji = '🔴' if priority == 'CRITICAL' else '🟠' if priority == 'HIGH' else '🟡' if priority == 'MEDIUM' else '🟢'
            doc.add_heading(f'{emoji} {priority} PRIORITY ({len(priority_tasks)} items)', 2)
            
            # Create table
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Task'
            header_cells[1].text = 'Owner'
            header_cells[2].text = 'Due Date'
            header_cells[3].text = 'Meeting'
            header_cells[4].text = 'Date'
            
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Add tasks
            for action in priority_tasks:
                row_cells = table.add_row().cells
                row_cells[0].text = action['what']
                row_cells[1].text = action['who']
                row_cells[2].text = str(action['due']) if action['due'] else 'TBD'
                row_cells[3].text = action['meeting']
                row_cells[4].text = action['date']
            
            doc.add_paragraph()
    
    # Tasks by owner
    doc.add_page_break()
    doc.add_heading('👥 Tasks by Owner', 1)
    
    owners = {}
    for action in all_actions:
        owner = action['who']
        if owner not in owners:
            owners[owner] = []
        owners[owner].append(action)
    
    for owner in sorted(owners.keys()):
        tasks = owners[owner]
        
        doc.add_heading(f'{owner} ({len(tasks)} tasks)', 2)
        
        # Create table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Task'
        header_cells[1].text = 'Due Date'
        header_cells[2].text = 'Priority'
        header_cells[3].text = 'Meeting'
        
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Add tasks
        for task in tasks:
            row_cells = table.add_row().cells
            row_cells[0].text = task['what']
            row_cells[1].text = str(task['due']) if task['due'] else 'TBD'
            row_cells[2].text = task['priority']
            row_cells[3].text = f"{task['meeting']} ({task['date']})"
        
        doc.add_paragraph()
    
    # Overdue tasks section
    doc.add_page_break()
    doc.add_heading('⚠️ Overdue & Urgent Tasks', 1)
    
    today_str = today.strftime('%Y-%m-%d')
    overdue_tasks = [a for a in all_actions if a['due'] and a['due'] != 'TBD' and a['due'] < today_str]
    
    if overdue_tasks:
        doc.add_paragraph(f'Found {len(overdue_tasks)} overdue task(s):')
        doc.add_paragraph()
        
        # Create table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Task'
        header_cells[1].text = 'Owner'
        header_cells[2].text = 'Due Date'
        header_cells[3].text = 'Days Overdue'
        header_cells[4].text = 'Priority'
        
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Add overdue tasks
        for task in overdue_tasks:
            row_cells = table.add_row().cells
            row_cells[0].text = task['what']
            row_cells[1].text = task['who']
            row_cells[2].text = task['due']
            
            # Calculate days overdue
            try:
                due_date = datetime.strptime(task['due'], '%Y-%m-%d').date()
                days_overdue = (today - due_date).days
                row_cells[3].text = str(days_overdue)
            except:
                row_cells[3].text = 'N/A'
            
            row_cells[4].text = task['priority']
    else:
        doc.add_paragraph('No overdue tasks! 🎉')
    
    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.add_run('Generated by 8825 Meeting Summary Pipeline').italic = True
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save document
    output_dir = pipeline.output_dir
    filename = f'Complete_Task_List_{datetime.now().strftime("%Y%m%d_%H%M")}.docx'
    filepath = output_dir / filename
    doc.save(filepath)
    
    print(f"\n{'='*80}")
    print(f"✅ Task list exported to Word document")
    print(f"{'='*80}\n")
    print(f"File: {filepath}")
    print(f"Total tasks: {len(all_actions)}")
    print(f"Overdue tasks: {len(overdue_tasks)}")
    print(f"\n{'='*80}\n")

if __name__ == "__main__":
    main()
